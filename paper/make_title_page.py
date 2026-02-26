#!/usr/bin/env python3
"""Generate the title page as a Word document."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('From Rules to Nash Equilibria: Formally Verified\nGame-Theoretic Analysis of a Competitive\nTrading Card Game')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(24)

# Authors
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Arthur F. Ramos')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
p.add_run('\n').font.size = Pt(13)
run2 = p.add_run('Microsoft, Redmond, WA, USA')
run2.font.size = Pt(12)
run2.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tulio Soria')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
p.add_run('\n').font.size = Pt(13)
run2 = p.add_run('Independent Researcher')
run2.font.size = Pt(12)
run2.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(24)

# Correspondence
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Corresponding Author:')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Arthur F. Ramos\nMicrosoft\nRedmond, WA, USA\n10158 Glenmore Ave, Bradenton, FL 34202\nEmail: arfreita@microsoft.com')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tulio Soria\nEmail: tulio.soria@gmail.com')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(24)

# Acknowledgments
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Acknowledgments')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'The authors thank the Pokémon TCG competitive community and the Trainer Hill '
    'platform for providing the tournament data used in this study. All formal '
    'verification was conducted using Lean 4. The authors declare no conflicts of interest.'
)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

output = '/Users/arthur/clawd/PokemonLean/paper/title_page.docx'
doc.save(output)
print(f"Saved: {output}")
