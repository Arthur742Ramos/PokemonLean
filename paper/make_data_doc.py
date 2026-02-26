#!/usr/bin/env python3
"""Generate a PDF data documentation file for IEEE DataPort."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p

# ============================================================
heading('Data Documentation: Pokémon TCG Tournament Matchup Data\nand Formally Verified Game-Theoretic Analysis')

para('This document describes the structure, provenance, and interpretation of all files included in the dataset.')

# ============================================================
heading('1. Data Source and Collection', level=2)

para('All empirical data were scraped from Trainer Hill (https://www.trainerhill.com/meta?game=PTCG) on February 19, 2026. The query filters were:')
bullet('Game: Pokémon Trading Card Game (PTCG)')
bullet('Tournament size: 50+ players')
bullet('Date range: January 29 – February 19, 2026')
bullet('Platforms: all')
bullet('Tie counting: 1/3 win (i.e., WR% = (W + T/3) / (W + L + T))')

para('Trainer Hill aggregates results from officially sanctioned Pokémon TCG tournaments. The data represent thousands of matches across dozens of events during the Stellar Crown – Prismatic Evolutions (SVI-PFL) format.')

# ============================================================
heading('2. File Descriptions', level=2)

heading('2.1 trainerhill_matchups_2026-02-19.md', level=3)
para('Format: Markdown with tables')
para('Size: ~12 KB (325 lines)')
para('This is the primary empirical data file containing:', bold=True)
bullet('Meta shares (overall): popularity of each of the 14 archetypes across all tournaments')
bullet('Top-8 shares: popularity among top-8 finishers only')
bullet('Full 14×14 matchup matrix: for each ordered pair of decks (row = player, column = opponent), the file reports the win rate percentage and the raw win-loss-tie record')
bullet('Condensed 6×6 matrix: a subset of the top 6 decks for quick reference')
bullet('Sample sizes: total match counts for key matchups (e.g., Dragapult mirror: 2,845 games)')

heading('The 14 Archetypes', level=3)
table = doc.add_table(rows=15, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['#', 'Archetype', 'Meta Share (%)']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
decks = [
    ('1', 'Dragapult Dusknoir', '15.5'),
    ('2', 'Gholdengo Lunatone', '9.9'),
    ('3', 'Grimmsnarl Froslass', '5.1'),
    ('4', 'Mega Absol Box', '5.0'),
    ('5', 'Other (aggregated)', '4.8'),
    ('6', 'Gardevoir', '4.6'),
    ('7', 'Charizard Noctowl', '4.3'),
    ('8', 'Gardevoir Jellicent', '4.2'),
    ('9', 'Dragapult Charizard', '3.5'),
    ('10', 'Raging Bolt Ogerpon', '3.4'),
    ('11', 'Alakazam Dudunsparce', '3.2'),
    ('12', "N's Zoroark", '3.1'),
    ('13', 'Ceruledge', '2.8'),
    ('14', 'Kangaskhan Bouffalant', '2.5'),
]
for i, (num, name, share) in enumerate(decks):
    table.rows[i+1].cells[0].text = num
    table.rows[i+1].cells[1].text = name
    table.rows[i+1].cells[2].text = share

heading('Reading the Matchup Matrix', level=3)
para('Each archetype section contains a table with columns:')
bullet('vs — the opponent archetype')
bullet('WR% — win rate as a percentage, computed as (W + T/3) / (W + L + T)')
bullet('W-L-T — raw win, loss, and tie counts')
para('The matrix is not perfectly symmetric because WR_A_vs_B + WR_B_vs_A ≈ 100% but not exactly, due to the tie-counting formula.')

# ============================================================
heading('2.2 bootstrap_nash_results.txt', level=3)
para('Format: Plain text')
para('Size: ~4 KB (90 lines)')
para('Contains the results of a 10,000-iteration bootstrap stability analysis of the Nash equilibrium. For each iteration, the 14×14 matchup matrix was resampled (with replacement from original W-L-T counts), and a new Nash equilibrium was computed. The file reports:')
bullet('Support frequency: how often each deck appears in the Nash support across 10,000 resamples')
bullet('Weight statistics: mean, median, standard deviation, and Wilson 95% confidence intervals for each deck\'s Nash weight')
bullet('Game value: the expected payoff at equilibrium (47.96% ± CI)')
bullet('Core support decks appear in >96% of resampled equilibria, confirming qualitative stability')

heading('2.3 symmetric_nash_results.txt', level=3)
para('Format: Plain text')
para('Size: ~4 KB (77 lines)')
para('Contains the Nash equilibrium of the symmetrized constant-sum game. The symmetrization transforms the matchup matrix M into M\' where M\'_ij = (M_ij + (1000 - M_ji)) / 2, yielding a zero-sum game. The file reports:')
bullet('Maximum deviation from perfect symmetry: max|M_ij + M_ji - 1000| = 73 permil')
bullet('Symmetric Nash strategy: 7-deck support with exact rational weights')
bullet('Game value: exactly 500 permil (as expected for a symmetric zero-sum game)')
bullet('Dragapult Dusknoir receives 0% weight despite 15.5% meta share (the "popularity paradox")')

heading('2.4 predictive_validation.txt', level=3)
para('Format: Plain text')
para('Size: ~4 KB (68 lines)')
para('Contains a predictive validation test of the replicator dynamics model. Single-step replicator dynamics on the full 14-deck game were computed using meta shares as initial population frequencies and the matchup matrix as the payoff matrix. Predictions (grow/shrink/stable) were compared against actual Trainer Hill trend arrows observed between February 19 and 20, 2026. The file reports:')
bullet('Per-deck fitness values and predicted direction')
bullet('Observed trend arrows from Trainer Hill')
bullet('Match/mismatch classification: 2 out of 3 testable predictions confirmed')

heading('2.5 native_decide_audit.txt', level=3)
para('Format: Plain text')
para('Size: ~7 KB (137 lines)')
para('An audit of all native_decide proof closures in the Lean 4 artifact. native_decide is a Lean tactic that trusts the compiler\'s evaluation rather than constructing kernel-checkable proof terms. This file catalogues every use, enabling reviewers to assess the trust boundary. The file reports:')
bullet('File-by-file listing of every native_decide invocation')
bullet('What each invocation proves (e.g., matrix entry verification, Nash optimality check)')
bullet('Total count and classification by proof category')

# ============================================================
heading('3. Units and Conventions', level=2)

bullet('Win rates in the Markdown file are percentages (0–100). Win rates in the Lean formalization use permil (0–1000) for exact integer arithmetic.')
bullet('Meta shares are percentages of total tournament entries.')
bullet('Ties are counted as 1/3 of a win per Trainer Hill convention.')
bullet('Nash equilibrium weights are reported as percentages (summing to 100%).')
bullet('The "Other" category aggregates all archetypes outside the top 14.')

# ============================================================
heading('4. Relationship to the Lean 4 Artifact', level=2)

para('The empirical data in this dataset are encoded in the Lean 4 proof artifact (available separately) in two canonical representations:')
bullet('RealMetagame.lean — matchup win rates as a function over the Deck inductive type (permil)')
bullet('NashEquilibrium.lean — matchup data as an Array (Array Nat) for computational proofs')
para('A machine-checked proof (MatrixConsistency.lean) establishes that both representations agree entry-by-entry. All game-theoretic results in the accompanying paper trace to this single empirical source.')

# ============================================================
heading('5. Reproducibility', level=2)

para('To reproduce the computed results from the raw data:')
bullet('Nash equilibria: use any linear programming solver (e.g., scipy.optimize.linprog) on the 14×14 matrix')
bullet('Bootstrap analysis: resample W-L-T counts with replacement, recompute Nash per iteration')
bullet('Replicator dynamics: apply the standard replicator equation dx_i/dt = x_i(f_i - f̄) with the matchup matrix as payoff')
para('Python scripts for all three computations are included in the accompanying code artifact.')

# ============================================================
heading('6. Citation', level=2)

para('If you use this dataset, please cite:')
para('A. F. Ramos and T. Soria, "From Rules to Nash Equilibria: Formally Verified Game-Theoretic Analysis of a Competitive Trading Card Game," submitted to IEEE Transactions on Games, 2026.')

output = '/Users/arthur/clawd/PokemonLean/paper/data_documentation.docx'
doc.save(output)
print(f"Saved: {output}")
